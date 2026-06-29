"""
highlight_word.py
根据实验完成情况，对老师建议文档添加高亮：
  - 已完成实验：黄色高亮
  - 未完成实验：红色高亮
"""
import shutil
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from lxml import etree

SRC = "新建 Microsoft Word 文档 (13).docx"
DST = "实验补充建议_高亮版.docx"

shutil.copy2(SRC, DST)
doc = Document(DST)

# ── 高亮颜色常量 ──────────────────────────────────────────
YELLOW = WD_COLOR_INDEX.YELLOW   # 已完成
RED    = WD_COLOR_INDEX.RED      # 未完成


def highlight_paragraph(para, color):
    """对整段文字施加高亮（保留原有格式）。"""
    for run in para.runs:
        run.font.highlight_color = color
    # 若段落无 run（只有段落级文字），补一个 run
    if not para.runs and para.text.strip():
        run = para.add_run()
        run.text = ""            # 内容已在 XML 中，不重复
        run.font.highlight_color = color


def highlight_table(table, color):
    """对整张表格所有单元格施加高亮。"""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                highlight_paragraph(para, color)


# ── 确定每个段落所属章节 ──────────────────────────────────
# 规则：遇到章节标题时更新 current_color；
#       '一、' / '二、' / '三、' / '4. 论文贡献' 是容器标题，不加高亮。
SECTION_RULES = [
    # (关键词, 颜色)
    ("2.1", YELLOW),
    ("2.2", YELLOW),
    ("2.3", YELLOW),
    ("2.4", YELLOW),
    ("2.5", RED),
    ("2.6", YELLOW),
    ("2.7", YELLOW),
    ("2.8", YELLOW),
    ("2.9", RED),
    ("2.10", RED),
    ("2.11", RED),
    ("2.12", RED),
    ("3. 推荐的 TBIOM", None),      # 推荐包标题，单独处理
    ("必补 1", YELLOW),
    ("必补 2", YELLOW),
    ("必补 3", YELLOW),
    ("必补 4", YELLOW),
    ("必补 5", YELLOW),
    ("必补 6", YELLOW),
    ("4. 论文贡献", None),           # 写作建议，不高亮
]

CONTAINER_KEYWORDS = ["一、必须补", "二、强烈建议", "三、加分项"]

def detect_section(text):
    """返回 (颜色 or None, is_container)"""
    for kw in CONTAINER_KEYWORDS:
        if kw in text:
            return None, True
    for kw, color in SECTION_RULES:
        if text.strip().startswith(kw):
            return color, False
    return "inherit", False   # "inherit" 代表沿用上一节颜色


# ── 遍历文档 body 中的所有元素（段落 + 表格） ──────────────────
body = doc.element.body
current_color = None   # 当前章节颜色

# 先收集 body 中所有段落和表格，按顺序处理
elements = []
for child in body:
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        # 找对应的 Paragraph 对象
        for para in doc.paragraphs:
            if para._element is child:
                elements.append(('para', para))
                break
    elif tag == 'tbl':
        # 找对应的 Table 对象
        for table in doc.tables:
            if table._element is child:
                elements.append(('table', table))
                break

for kind, obj in elements:
    if kind == 'para':
        para = obj
        text = para.text.strip()
        if not text:
            continue

        color, is_container = detect_section(text)

        if is_container:
            current_color = None      # 容器标题不高亮，下面的正文沿用各自节颜色
            continue

        if color == "inherit":
            # 正文，沿用当前节颜色
            if current_color is not None:
                highlight_paragraph(para, current_color)
        elif color is None:
            current_color = None      # 论文贡献部分不高亮
        else:
            current_color = color
            highlight_paragraph(para, current_color)

    elif kind == 'table':
        if current_color is not None:
            highlight_table(obj, current_color)

doc.save(DST)
print(f"已保存：{DST}")
